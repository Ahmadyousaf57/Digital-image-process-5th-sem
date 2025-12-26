# save file as make_architecture_diagram.py and run: python make_architecture_diagram.py
from graphviz import Digraph
from docx import Document
from docx.shared import Inches

# ---------- Architecture Flow Diagram (styled) ----------
arch = Digraph('ArchitectureFlow', format='png')
arch.attr(rankdir='LR', size='10', bgcolor='white')

# Nodes with nicer shapes/colors
arch.node('Camera', label='Camera\n(Video / Image Capture)', shape='parallelogram', style='filled', fillcolor='#BEE3F8') 
arch.node('Processing', label='Face Recognition\nProcessing Unit\n(DIP + Model)', shape='oval', style='filled', fillcolor='#C6F6D5')
arch.node('DB', label='Authorized User\nDatabase\n(Templates + Logs)', shape='cylinder', style='filled', fillcolor='#FEF3C7')
arch.node('Access', label='Access Control\n(Gate / Lock)', shape='box', style='filled', fillcolor='#FED7E2')
arch.node('Admin', label='Admin Dashboard\n(Monitor / Alerts)', shape='box3d', style='filled', fillcolor='#E9D8FD')

# Edges / data flows with labels
arch.edge('Camera', 'Processing', label='Captured Image', fontsize='10')
arch.edge('Processing', 'DB', label='Match Request / Template Update', fontsize='10')
arch.edge('DB', 'Processing', label='Template / User Data', fontsize='10')
arch.edge('Processing', 'Access', label='Allow / Deny Signal', fontsize='10')
arch.edge('Processing', 'Admin', label='Logs & Alerts', fontsize='10', style='dashed')

# Optional: cluster to show system boundary
arch.attr('node', shape='none')
arch.node('boundary', label='<<TABLE BORDER="0" CELLBORDER="0" CELLPADDING="4">'
                           '<TR><TD ALIGN="LEFT"><B>Face Recognition System Boundary</B></TD></TR></TABLE>>')
# place the boundary (visual only) - not strict positioning but shows a title
# render the diagram
arch.render('Architecture_Flow_Diagram', cleanup=True)
print("Saved Architecture_Flow_Diagram.png")

# ---------- Insert the diagram into a Word file ----------
doc = Document()
doc.add_heading("Architecture Flow", level=1)
doc.add_paragraph("Flow diagram for the Face Recognition based Entry System:")
doc.add_picture("Architecture_Flow_Diagram.png", width=Inches(6))
doc.add_paragraph("\nLegend:\n- Camera: captures image/video\n- Processing: DIP + recognition model\n- Database: stores templates and logs\n- Access Control: gate/lock which receives allow/deny signal\n- Admin Dashboard: monitoring and alerts")
doc.save("Architecture_Report_With_Diagram.docx")
print("Saved Architecture_Report_With_Diagram.docx")
